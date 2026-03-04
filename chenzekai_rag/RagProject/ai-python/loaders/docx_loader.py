"""DOCX loader implementation."""

from __future__ import annotations

from pathlib import Path

from .base import BaseLoader, ParsedDocument


class DocxLoader(BaseLoader):
    """Load DOCX documents into ParsedDocument objects."""

    def load(self, path: str) -> ParsedDocument:
        """Parse a DOCX file and return a ParsedDocument."""
        try:
            import docx  # type: ignore
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise RuntimeError("python-docx is required for DOCX loading") from exc

        document = docx.Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        source = Path(path)
        return ParsedDocument(
            doc_id=source.stem,
            text=text.strip(),
            metadata={"source_type": "docx"},
            source_path=str(source),
        )
